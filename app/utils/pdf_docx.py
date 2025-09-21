import os
from io import BytesIO
from fpdf import FPDF
from docx import Document
from fastapi.responses import StreamingResponse

def mcqs_to_pdf(mcqs: list) -> StreamingResponse:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for mcq in mcqs:
        pdf.multi_cell(0, 10, f"Q{mcq['number']}: {mcq['question']}")
        for idx, opt in enumerate(mcq['options']):
            letter = chr(65 + idx)  # A, B, C, D
            pdf.multi_cell(0, 10, f"   {letter}) {opt}")
        pdf.multi_cell(0, 10, f"Correct Answer: {mcq['correct_answer']}")
        pdf.multi_cell(0, 10, "")
    pdf_bytes = pdf.output(dest="S").encode("latin-1", "replace")
    return StreamingResponse(BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=mcqs_output.pdf"}
    )

def mcqs_to_docx(mcqs: list) -> StreamingResponse:
    doc = Document()
    doc.add_heading("Generated MCQs", level=1)
    for mcq in mcqs:
        doc.add_paragraph(f"Q{mcq['number']}: {mcq['question']}")
        for idx, opt in enumerate(mcq['options']):
            letter = chr(65 + idx)  # A, B, C, D
            doc.add_paragraph(f"   {letter}) {opt}")
        doc.add_paragraph(f"Correct Answer: {mcq['correct_answer']}")
        doc.add_paragraph("")
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return StreamingResponse(out,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=mcqs_output.docx"}
    )